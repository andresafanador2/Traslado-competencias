import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
import os

# =====================================
# JSON / Diccionarios para los tres tipos de traslado
# =====================================

# Diccionario con municipios y correos (Traslado por cobro AP)
municipios_dict = {
    "Abrego": {
    "correos": [
      "secretariadehacienda@abregonortedesantander.gov.co",
      "secretariadegobierno@abrego-nortedesantander.gov.co"
    ]
  },
  "Aguachica": {
    "correos": [
      "alcaldia@aguachica-cesar.gov.co",
      "hacienda@aguachica-cesar.gov.co"
    ]
  },
  "Arboledas": {
    "correos": [
      "alcaldia@arboledas-nortedesantander.gov.co",
      "usp@arboledas-nortedesantander.gov.co"
    ]
  },
  "Bochalema": {
    "correos": [
      "alcaldia@bochalema-nortedesantander.gov.co",
      "contactenos@bochalema-nortedesantander.gov.co"
    ]
  },
  "Bucarasica": {
    "correos": [
      "alcaldia@bucarasica-nortedesantander.gov.co",
      "gobierno@bucarasica-nortedesantander.gov.co"
    ]
  },
  "Cáchira": {
    "correos": [
      "alcaldia@cachira-nortedesantander.gov.co",
      "hacienda@cachira-nortedesantander.gov.co"
    ]
  },
  "Cácota": {
    "correos": [
      "alcaldia@cacota-nortedesantander.gov.co",
      "hacienda@cacota-nortedesantander.gov.co"
    ]
  },
  "Chinácota": {
    "correos": [
      "alcaldia@chinacota-nortedesantander.gov.co",
      "planeacion@chinacota-nortedesantander.gov.co"
    ]
  },
  "Chitagá": {
    "correos": [
      "alcaldia@chitaga-nortedesantander.gov.co",
      "planeacion@chitaga-nortedesantander.gov.co"
    ]
  },
  "Convención": {
    "correos": [
      "sechacienda@convencion-nortedesantander.gov.co",
      "SECPLANEACION@convencion-nortedesantander.gov.co"
    ]
  },
  "Cucutilla": {
    "correos": [
      "alcalde@cucutilla-nortedesantander.gov.co",
      "emiliagudelo092@gmail.com"
    ]
  },
  "Durania": {
    "correos": [
      "alcaldia@durania-nortedesantander.gov.co",
      "secretariageneral@durania-nortedesantander.gov.co"
    ]
  },
  "El Carmen": {
    "correos": [
      "alcaldia@elcarmen-nortedesantander.gov.co",
      "secretariadeplaneacion@elcarmen-nortedesantander.gov.co"
    ]
  },
  "El Tarra": {
    "correos": [
      "despacho@eltarra-nortedesantander.gov.co",
      "sechacienda@eltarra-nortedesantander.gov.co"
    ]
  },
  "El Zulia": {
    "correos": [
      "alcaldia@elzulia-nortedesantander.gov.co",
      "secretariadegobierno@elzulia-nortedesantander.gov.co"
    ]
  },
  "Gamarra": {
    "correos": [
      "alcaldia@gamarra-cesar.gov.co",
      "gobierno@gamarra-cesar.gov.co"
    ]
  },
  "González": {
    "correos": [
      "alcaldia@gonzalez-cesar.gov.com",
      "secretariadehacienda@gonzalezcesar.gov.co"
    ]
  },
  "Gramalote": {
    "correos": [
      "alcalde@gramalote-nortedesantander.gov.co",
      "planeacion@gramalote-nortedesantander.gov.co"
    ]
  },
  "Hacarí": {
    "correos": [
      "alcaldia@hacari-nortedesantander.gov.co",
      "secgobierno@hacari-nortedesantander.gov.co"
    ]
  },
  "Herrán": {
    "correos": [
      "alcaldia@herran-nortedesantander.gov.co",
      "secretariadehacienda@herran-nortedesantander.gov.co"
    ]
  },
  "La Esperanza": {
    "correos": [
      "alcalde@laesperanza-nortedesantander.gov.co",
      "hacienda@laesperanza-nortedesantander.gov.co"
    ]
  },
  "La Gloria": {
    "correos": [
      "alcaldia@lagloria-cesar.gov.co",
      "secretariadegobierno@lagloria-cesar.gov.co"
    ]
  },
  "La Playa": {
    "correos": [
      "despachoalcaldia@laplayadebelen-nortedesantander.gov.co",
      "haciendaysalud@laplayadebelen-nortedesantander.gov.co"
    ]
  },
  "Labateca": {
    "correos": [
      "alcaldia@labateca-nortedesantander.gov.co",
      "secretariadehacienda@labateca-nortedesantander.gov.co"
    ]
  },
  "Los Patios": {
    "correos": [
      "secretariadehacienda@lospatios-nortedesantander.gov.co",
      "ricardolospatios.planeacion@gmail.com"
    ]
  },
  "Lourdes": {
    "correos": [
      "alcaldia@lourdes-nortedesantander.gov.co",
      "tesoreria@lourdes-nortedesantander.gov.co"
    ]
  },
  "Morales": {
    "correos": [
      "hladimirdelgado337@gmail.com",
      "andriussierra@gmail.com"
    ]
  },
  "Mutiscua": {
    "correos": [
      "alcaldia@mutiscua-nortedesantander.gov.co",
      "tesoreria@mutiscua-nortedesantander.gov.co"
    ]
  },
  "Ocaña": {
    "correos": [
      "alcalde@ocana-nortedesantander.gov.co",
      "secretariadehacienda@ocana-nortedesantander.gov.co"
    ]
  },
  "Pamplona": {
    "correos": [
      "secretariadegobierno@pamplona-nortedesantander.gov.co",
      "secretariadeplaneacion@pamplona-nortedesantander.gov.co"
    ]
  },
  "Pamplonita": {
    "correos": [
      "alcaldia@pamplonita-nortedesantander.gov.co",
      "Tesoreria@pamplonita-nortedesantander.gov.co"
    ]
  },
  "Pelaya": {
    "correos": [
      "gobierno@pelaya-cesar.gov.co",
      "hacienda@pelaya-cesar.gov.co"
    ]
  },
  "Puerto Santander": {
    "correos": [
      "alcaldia@puertosantander-nortedesantander.gov.co",
      "planeacion@puertosantander-nortedesantander.gov.co"
    ]
  },
  "Ragonvalía": {
    "correos": [
      "alcaldia@ragonvalia-nortedesantander.gov.co",
      "—"
    ]
  },
  "Río de Oro": {
    "correos": [
      "alcaldia@riodeoro-cesar.gov.co",
      "sechacienda@riodeoro-cesar.gov.co"
    ]
  },
  "Salazar": {
    "correos": [
      "alcaldia@salazardelaspalmas-nortedesantander.gov.co",
      "hacienda@salazardelaspalmas-nortedesantander.gov.co"
    ]
  },
  "San Calixto": {
    "correos": [
      "alcaldia@sancalixto-nortedesantander.gov.co",
      "secretariadehacienda@sancalixtonortedesantander.gov.co"
    ]
  },
  "San Cayetano": {
    "correos": [
      "alcaldia@sancayetano-nortedesantander.gov.co",
      "planeacion@sancayetano-nortedesantander.gov.co"
    ]
  },
  "San José de Cúcuta": {
    "correos": [
      "impuestos@cucuta.gov.co",
      "planeacion@cucuta-nortedesantander.gov.co"
    ]
  },
  "Santiago": {
    "correos": [
      "alcaldia@santiago-nortedesantander.gov.co",
      "tesoreria@santiago-nortedesantander.gov.co"
    ]
  },
  "Sardinata": {
    "correos": [
      "alcaldia@sardinata-nortedesantander.gov.co",
      "planeacion@sardinata-nortedesantander.gov.co"
    ]
  },
  "Silos": {
    "correos": [
      "alcaldia@silos-nortedesantander.gov.co",
      "hacienda@silos-nortedesantander.gov.co"
    ]
  },
  "Teorama": {
    "correos": [
      "alcaldia@teorama-nortedesantander.gov.co",
      "planeacion@teorama-nortedesantander.gov.co"
    ]
  },
  "Tibú": {
    "correos": [
      "alcaldia@tibu-nortedesantander.gov.co",
      "secretariadeplaneacion@tibu-nortedesantander.gov.co"
    ]
  },
  "Toledo": {
    "correos": [
      "alcaldia@toledo-nortedesantander.gov.co",
      "hacienda@toledo-nortedesantander.gov.co"
    ]
  },
  "Villa Caro": {
    "correos": [
      "alcaldia@villacaro-nortedesantander.gov.co",
      "planeacion@villacaro-nortedesantander.gov.co"
    ]
  },
  "Villa del Rosario": {
    "correos": [
      "alcaldia@villadelrosario-nortedesantander.gov.co",
      "secretariadeplaneacion@villadelrosario-nortedesantander.gov.co"
    ]
  }
}

# JSON de mantenimiento AP
mantenimiento_ap = [
  {
    "responsable": "Energizett",
    "municipios": [
      {"municipio": "Los Patios", "correos": ["director.patios@energizett.com", "administrativo@energizett.com"]},
      {"municipio": "Bochalema", "correos": ["director.patios@energizett.com", "administrativo@energizett.com"]},
      {"municipio": "Chinácota", "correos": ["director.patios@energizett.com", "administrativo@energizett.com"]},
      {"municipio": "Chitagá", "correos": ["director.patios@energizett.com", "administrativo@energizett.com"]},
      {"municipio": "Pamplonita", "correos": ["director.patios@energizett.com", "administrativo@energizett.com"]},
      {"municipio": "Mutiscua", "correos": ["contactenos.mutiscua@energizett.com"]}
    ]
  },
  {
    "responsable": "Concesión Alumbrado Público SJC",
    "municipios": [
      {"municipio": "Abrego", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "Convención", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "Ocaña", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "Aguachica", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "Cúcuta", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "Gamarra", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "La Gloria", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "Pamplona", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "Puerto Santander", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "San Cayetano", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "Sardinata", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "Tibú", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]},
      {"municipio": "Villa del Rosario", "correos": ["atencionalusuario@alumbradosjc.com", "inginfraestructura@alumbradosjc.com"]}
    ]
  },
  {
    "responsable": "EMSAP S.A.S",
    "municipios": [
      {"municipio": "Santiago", "correos": ["gerencia@emsap.gov.co"]}
    ]
  },
  {
    "responsable": "EMTSAS",
    "municipios": [
      {"municipio": "El Tarra", "correos": ["gerencia@emtsas.gov.co"]}
    ]
  },
  {
    "responsable": "KERING S.A.S",
    "municipios": [
      {"municipio": "El Carmen", "correos": ["keringsas@gmail.com", "hacienda.elcarmen.iap@gmail.com"]},
      {"municipio": "Morales", "correos": ["keringsas@gmail.com", "hacienda.elcarmen.iap@gmail.com"]}
    ]
  },
  {
    "responsable": "Consorcio Lumen de Colombia",
    "municipios": [
      {"municipio": "El Zulia", "correos": ["j.jbuitrago12@gmail.com"]}
    ]
  },
  {
    "responsable": "UNION TEMPORAL ALUMBRADO DE PELAYA",
    "municipios": [
      {"municipio": "Pelaya", "correos": ["alumbrado.pelaya@gmail.com"]}
    ]
  },
  {
    "responsable": "DEYCON S.A.S",
    "municipios": [
      {"municipio": "Salazar", "correos": ["prydeyconsas@gmail.com"]}
    ]
  },
  {
    "responsable": "Empresa de Iluminación Pública y Tecnológica de Sardinata - IPSA S.A.S.",
    "municipios": [
      {"municipio": "Sardinata", "correos": ["alumbradosardinata@ipsa.com.co"]}
    ]
  }
]

# JSON de otras entidades
otras_entidades = 
[
    {
        "empresa": "VEOLIA",
        "correos": ["co.servicioalcliente.aseo.oriente@veolia.com"]
    },
    {
        "empresa": "AGUAS KAPITAL",
        "correos": ["radicacion.ceindoc@akc.co"]
    },
    {
        "empresa": "AFINIA",
        "correos": ["correspondencia@afinia.com.co"]
    },
    {
        "empresa": "Empresas Públicas de Medellin",
        "correos": ["epm@epm.com.co"]
    },
    {
        "empresa": "ESSA",
        "correos": ["essa@essa.com.co"]
    },
    {
        "empresa": "ENEL",
        "correos": ["servicioalcliente.empresarial@enel.com"]
    },
    {
        "empresa": "CENTRALES ELECTRICAS DE NARIÑO - CEDENAR",
        "correos": ["correspondenciacad@cedenar.com.co"]
    },
    {
        "empresa": "VATIA",
        "correos": ["linsac@vatia.com.co"]
    },
    {
        "empresa": "Electrificadora del Huila S.A. E.S.P.",
        "correos": ["radicacion@electrohuila.co"]
    },
    {
        "empresa": "Centrales Eléctricas del Cauca S.A. E.S.P.",
        "correos": ["contacto@cedelca.com.co"]
    },
    {
        "empresa": "Empresa de Energía de Pereira S.A. E.S.P.",
        "correos": ["contactenos@eep.com.co"]
    },
    {
        "empresa": "Electrificadora del Meta S.A. E.S.P.",
        "correos": ["pqr@emsa-esp.com.co"]
    },
    {
        "empresa": "Empresa Distribuidora del Pacífico S.A. E.S.P.",
        "correos": ["acliente@dispacsaesp.com"]
    },
    {
        "empresa": "Empresa de Energía del Quindío S.A. E.S.P.",
        "correos": ["edeq@edeq.com.co"]
    }
]

# Lista de municipios de Norte de Santander
municipios_norte_santander = [
    "Arboledas", "Cucutilla", "Gramalote", "Lourdes", "Salazar de Las Palmas",
    "Santiago", "Villa Caro", "Cúcuta", "El Zulia", "Los Patios",
    "Puerto Santander", "San Cayetano", "Villa del Rosario", "Bucarasica",
    "El Tarra", "Sardinata", "Tibú", "Ábrego", "Cáchira", "Convención",
    "El Carmen", "Hacarí", "La Esperanza", "La Playa de Belén", "Ocaña",
    "San Calixto", "Teorama", "Cácota", "Chitagá", "Mutiscua",
    "Pamplona", "Pamplonita", "Santo Domingo de Silos", "Bochalema",
    "Chinácota", "Durania", "Herrán", "Labateca", "Ragonvalia", "Toledo"
]

# =====================================
# Variables globales
# =====================================
responsable_seleccionado = ""

# =====================================
# Crear ventana principal
# =====================================
root = tk.Tk()
root.title("Gestión de Traslados")
root.geometry("700x600")

tk.Label(root, text="Bienvenido al Sistema de Trámites", font=("Arial", 16, "bold")).pack(pady=10)
tk.Label(root, text="Seleccione el tipo de traslado que desea realizar:", font=("Arial", 12)).pack(pady=5)

opciones = ["Traslado por cobro AP", "Traslado por mantenimiento AP", "Traslado otras entidades"]
opcion_cb = ttk.Combobox(root, values=opciones, state="readonly", width=50)
opcion_cb.pack(pady=10)

frame_formulario = tk.Frame(root)
frame_formulario.pack(pady=10, fill="both", expand=True)

# =====================================
# Función para reemplazar texto manteniendo fuente y tamaño
# =====================================
def reemplazar_texto_parrafos(doc, reemplazos: dict):
    for p in doc.paragraphs:
        for key, valor in reemplazos.items():
            if key in p.text:
                for run in p.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, valor)

# =====================================
# Función para generar documentos
# =====================================
def generar_documentos():
    global responsable_seleccionado
    try:
        # Datos del formulario
        municipio = municipio_cb.get()
        nombre_usuario = entry_nombre.get()
        direccion = entry_direccion.get()
        correo_usuario = entry_correo.get()
        radicado = entry_radicado.get()
        fecha = entry_fecha.get()
        expediente = entry_expediente.get()
        proceso = entry_proceso.get()

        if not municipio or not nombre_usuario:
            messagebox.showerror("Error", "Debe diligenciar al menos municipio y nombre de usuario.")
            return

        # Rutas dinámicas
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        plantilla_traslado = os.path.join(desktop, "Modelo_Traslado.docx")
        plantilla_info_usuario = os.path.join(desktop, "MODELO_INFO_USUARIO_Modificada.docx")

        if not os.path.exists(plantilla_traslado) or not os.path.exists(plantilla_info_usuario):
            messagebox.showerror("Error", "No se encuentran las plantillas en el escritorio.")
            return

        # Determinar responsable y correos según tipo de traslado
        if opcion_cb.get() == "Traslado por cobro AP":
            responsable_texto = f"ALCALDIA {municipio.upper()}"
            correos = "; ".join(municipios_dict.get(municipio, {}).get("correos", []))
        elif opcion_cb.get() == "Traslado por mantenimiento AP":
            responsable_texto = responsable_seleccionado
            correos = ""
            for item in mantenimiento_ap:
                if item["responsable"] == responsable_seleccionado:
                    for m in item["municipios"]:
                        if m["municipio"] == municipio:
                            correos = "; ".join(m["correos"])
        else:  # Traslado otras entidades
            responsable_texto = responsable_seleccionado
            correos = ""
            for item in otras_entidades:
                if item["empresa"] == responsable_seleccionado:
                    correos = "; ".join(item["correos"])

        # Documento Modelo_Traslado
        doc1 = Document(plantilla_traslado)
        reemplazos1 = {
            "(RESPONSABLE)": responsable_texto,
            "(CORREO)": correos,
            "(RADICADO)": radicado,
            "(FECHA)": fecha,
            "(EXPEDIENTE)": expediente,
            "(PROCESO)": proceso
        }
        reemplazar_texto_parrafos(doc1, reemplazos1)
        salida1 = os.path.join(desktop, f"Traslado_{municipio}.docx")
        doc1.save(salida1)

        # Determinar departamento según municipio
        if municipio in municipios_norte_santander:
            departamento = "Norte de Santander"
        else:
            departamento = "Cesar"

        # Documento MODELO_INFO_USUARIO
        doc2 = Document(plantilla_info_usuario)
        reemplazos2 = {
            "(NOMBRE USUARIO)": nombre_usuario,
            "(DIRECCION)": direccion,
            "(CORREO)": correo_usuario,
            "(MUNICIPIO)": municipio,
            "(DEPARTAMENTO)": departamento,  # <-- cambio aplicado
            "(RADICADO)": radicado,
            "(FECHA)": fecha,
            "(EXPEDIENTE)": expediente,
            "(PROCESO)": proceso,
            "(RESPONSABLE)": responsable_texto
        }
        reemplazar_texto_parrafos(doc2, reemplazos2)
        salida2 = os.path.join(desktop, f"Info_Usuario_{nombre_usuario}.docx")
        doc2.save(salida2)

        messagebox.showinfo("Éxito", f"Se generaron los documentos en el escritorio:\n{salida1}\n{salida2}")

    except Exception as e:
        messagebox.showerror("Error", f"Ocurrió un error: {e}")

# =====================================
# Función para mostrar formulario según tipo
# =====================================
def mostrar_formulario(opcion):
    global municipio_cb, combo_responsable, entry_nombre, entry_direccion, entry_correo
    global entry_radicado, entry_fecha, entry_expediente, entry_proceso, responsable_seleccionado

    for widget in frame_formulario.winfo_children():
        widget.destroy()

    if opcion == "Traslado por cobro AP":
        tk.Label(frame_formulario, text="Municipio:").grid(row=0, column=0, sticky="w")
        municipio_cb = ttk.Combobox(frame_formulario, values=list(municipios_dict.keys()), state="readonly")
        municipio_cb.grid(row=0, column=1, sticky="w")
        responsable_seleccionado = ""  # Se definirá como la alcaldía al generar doc

        labels = ["Nombre del usuario", "Dirección", "Correo", "Radicado", "Expediente", "Proceso", "Fecha"]
        entries = []
        for i, lbl in enumerate(labels, start=1):
            tk.Label(frame_formulario, text=f"{lbl}:").grid(row=i, column=0, sticky="w")
            entry = tk.Entry(frame_formulario, width=50)
            entry.grid(row=i, column=1, sticky="w")
            entries.append(entry)
        entry_nombre, entry_direccion, entry_correo, entry_radicado, entry_expediente, entry_proceso, entry_fecha = entries

        tk.Button(frame_formulario, text="Generar Documentos", command=generar_documentos, bg="green", fg="white")\
            .grid(row=len(labels)+1, column=0, columnspan=2, pady=10)

    elif opcion == "Traslado por mantenimiento AP":
        tk.Label(frame_formulario, text="Responsable:").grid(row=0, column=0, sticky="w")
        responsables_lista = [item["responsable"] for item in mantenimiento_ap]
        combo_responsable = ttk.Combobox(frame_formulario, values=responsables_lista, state="readonly")
        combo_responsable.grid(row=0, column=1, sticky="w")

        tk.Label(frame_formulario, text="Municipio:").grid(row=1, column=0, sticky="w")
        municipio_cb = ttk.Combobox(frame_formulario, state="readonly")
        municipio_cb.grid(row=1, column=1, sticky="w")

        tk.Label(frame_formulario, text="Correos:").grid(row=2, column=0, sticky="w")
        correos_var = tk.StringVar()
        tk.Entry(frame_formulario, textvariable=correos_var, width=50, state="readonly").grid(row=2, column=1, sticky="w")

        labels = ["Nombre del usuario", "Dirección", "Correo", "Radicado", "Expediente", "Proceso", "Fecha"]
        entries = []
        for i, lbl in enumerate(labels, start=3):
            tk.Label(frame_formulario, text=f"{lbl}:").grid(row=i, column=0, sticky="w")
            entry = tk.Entry(frame_formulario, width=50)
            entry.grid(row=i, column=1, sticky="w")
            entries.append(entry)
        entry_nombre, entry_direccion, entry_correo, entry_radicado, entry_expediente, entry_proceso, entry_fecha = entries

        tk.Button(frame_formulario, text="Generar Documentos", command=generar_documentos, bg="green", fg="white")\
            .grid(row=len(labels)+3, column=0, columnspan=2, pady=10)

        def actualizar_responsable(event):
            global responsable_seleccionado
            sel = combo_responsable.get()
            responsable_seleccionado = sel
            for item in mantenimiento_ap:
                if item["responsable"] == sel:
                    municipio_cb["values"] = [m["municipio"] for m in item["municipios"]]
                    municipio_cb.set("")
                    correos_var.set("")

        combo_responsable.bind("<<ComboboxSelected>>", actualizar_responsable)

        def actualizar_municipio(event):
            mun_sel = municipio_cb.get()
            for item in mantenimiento_ap:
                if item["responsable"] == responsable_seleccionado:
                    for m in item["municipios"]:
                        if m["municipio"] == mun_sel:
                            correos_var.set("; ".join(m["correos"]))

        municipio_cb.bind("<<ComboboxSelected>>", actualizar_municipio)

    elif opcion == "Traslado otras entidades":
        tk.Label(frame_formulario, text="Empresa / Responsable:").grid(row=0, column=0, sticky="w")
        empresas_lista = [item["empresa"] for item in otras_entidades]
        combo_responsable = ttk.Combobox(frame_formulario, values=empresas_lista, state="readonly")
        combo_responsable.grid(row=0, column=1, sticky="w")

        tk.Label(frame_formulario, text="Correos asociados:").grid(row=1, column=0, sticky="w")
        correos_var = tk.StringVar()
        tk.Entry(frame_formulario, textvariable=correos_var, width=50, state="readonly").grid(row=1, column=1, sticky="w")

        labels = ["Nombre del usuario", "Dirección", "Correo", "Radicado", "Expediente", "Proceso", "Fecha", "Municipio"]
        entries = []
        for i, lbl in enumerate(labels, start=2):
            tk.Label(frame_formulario, text=f"{lbl}:").grid(row=i, column=0, sticky="w")
            entry = tk.Entry(frame_formulario, width=50)
            entry.grid(row=i, column=1, sticky="w")
            entries.append(entry)
        entry_nombre, entry_direccion, entry_correo, entry_radicado, entry_expediente, entry_proceso, entry_fecha, municipio_cb = entries

        tk.Button(frame_formulario, text="Generar Documentos", command=generar_documentos, bg="green", fg="white")\
            .grid(row=len(labels)+2, column=0, columnspan=2, pady=10)

        def actualizar_empresa(event):
            global responsable_seleccionado
            sel = combo_responsable.get()
            responsable_seleccionado = sel
            for item in otras_entidades:
                if item["empresa"] == sel:
                    correos_var.set("; ".join(item["correos"]))

        combo_responsable.bind("<<ComboboxSelected>>", actualizar_empresa)

# =====================================
# Bind del combo principal
# =====================================
opcion_cb.bind("<<ComboboxSelected>>", lambda event: mostrar_formulario(opcion_cb.get()))

root.mainloop()


